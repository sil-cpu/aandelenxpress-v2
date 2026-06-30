"""
Exporteert alle vragenlijst-HTML bestanden naar Word-documenten op het bureaublad.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches

PUBLIC_DIR = Path(__file__).parent.parent / "public"
DESKTOP    = Path.home() / "Desktop"

FILES = {
    "BV Oprichten":            "vragenlijst-bv-oprichten.html",
    "BV en Holding Oprichten": "vragenlijst-bv-holding.html",
    "Holding Oprichten":       "vragenlijst-holding-oprichten.html",
    "Geruisloze Inbreng":      "vragenlijst-geruisloos.html",
    "Eenmanszaak naar BV":     "vragenlijst-eenmanszaak-naar-bv.html",
    "VOF naar BV":             "vragenlijst-vof-naar-bv.html",
}

def strip_tags(s):
    s = re.sub(r'<[^>]+>', ' ', s)
    s = re.sub(r'&nbsp;', ' ', s)
    s = re.sub(r'&mdash;', '—', s)
    s = re.sub(r'&euro;', '€', s)
    s = re.sub(r'&amp;', '&', s)
    s = re.sub(r'&#[0-9]+;', '', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip(" *\n\r")


def extract_form(html):
    """
    Parses the vragenlijst HTML structure:
    <div class="card">
      <div class="card-title">Section</div>
      <div class="field">
        <label class="field-label">Question</label>
        <ul class="options-list">
          <li><div class="option-label">X</div><div class="option-desc">Y</div></li>
        </ul>
        OR <input> OR <textarea> OR <select>
        <div class="field-hint">...</div>
      </div>
    </div>
    """
    # Strip style/script
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.S)
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.S)

    sections = []

    # Split into card blocks: find each <div class="card"> opening and get content until next card
    # Use a non-greedy split based on card-header presence
    card_pattern = re.compile(
        r'<div[^>]*class="[^"]*\bcard\b[^"]*"[^>]*>(.*?)<div[^>]*class="[^"]*\bcard-header\b[^"]*"[^>]*>(.*?)</div>\s*<div[^>]*class="[^"]*\bcard-body\b[^"]*"[^>]*>(.*?)(?=<div[^>]*class="[^"]*\bcard\b[^"]*">|$)',
        re.S
    )

    # Simpler: extract card blocks by finding card-title + associated card-body
    # Step 1: find all card divs (split on card opening)
    # Split HTML on every <div class="card"> or <div class="card featured"> etc.
    raw_cards = re.split(r'(?=<div\s+class="card")', html)

    for raw in raw_cards:
        # Get title — support both card-title and card-h variants
        title_m = re.search(
            r'<div[^>]*class="[^"]*(?:card-title|card-h)[^"]*"[^>]*>(.*?)</div>',
            raw, re.S
        )
        if not title_m:
            continue
        title = strip_tags(title_m.group(1))
        if not title:
            continue

        fields = []

        # Get card body — support both card-body and card-b variants
        body_m = re.search(
            r'<div[^>]*class="[^"]*(?:card-body|card-b)[^"]*"[^>]*>(.*)',
            raw, re.S
        )
        body = body_m.group(1) if body_m else raw

        # Split body into field blocks
        # A field starts with <div class="field"...> or <div class="reveal"...><div class="field"...>
        # We split on the opening of field divs
        field_chunks = re.split(r'(?=<div[^>]*(?:class="[^"]*\bfield\b[^"]*"|id="field[A-Z][^"]*")[^>]*>)', body)

        for fc in field_chunks:
            # Must have a field-label or plain label inside .field
            lbl_m = (
                re.search(r'<label[^>]*class="[^"]*field-label[^"]*"[^>]*>(.*?)</label>', fc, re.S)
                or re.search(r'<label(?:\s[^>]*)?>([^<]{3,})</label>', fc, re.S)
            )
            if not lbl_m:
                continue
            label = strip_tags(lbl_m.group(1))
            if not label or label == '*':
                continue

            field = {
                "label": label,
                "type": "",
                "hint": "",
                "options": [],        # list of (label, desc) tuples
                "placeholder": "",
                "extra_text": "",
            }

            # Hint
            hint_m = re.search(r'<div[^>]*class="[^"]*field-hint[^"]*"[^>]*>(.*?)</div>', fc, re.S)
            if hint_m:
                field["hint"] = strip_tags(hint_m.group(1))

            # Extra descriptive text (p tags with info, not hints)
            ptexts = re.findall(r'<p[^>]*style="[^"]*font-size[^"]*"[^>]*>(.*?)</p>', fc, re.S)
            if ptexts:
                field["extra_text"] = strip_tags(ptexts[0])

            # Options list (radio / checkbox)
            if re.search(r'class="[^"]*options-list[^"]*"', fc):
                has_cb = bool(re.search(r'type="checkbox"', fc))
                field["type"] = "Meerkeuze" if has_cb else "Keuze (één antwoord)"

                # Extract option items
                option_items = re.findall(
                    r'<li[^>]*class="[^"]*option-item[^"]*"[^>]*>(.*?)</li>',
                    fc, re.S
                )
                for item in option_items:
                    olbl_m = re.search(r'<div[^>]*class="[^"]*option-label[^"]*"[^>]*>(.*?)</div>', item, re.S)
                    odsc_m = re.search(r'<div[^>]*class="[^"]*option-desc[^"]*"[^>]*>(.*?)</div>', item, re.S)
                    if olbl_m:
                        olbl = strip_tags(olbl_m.group(1))
                        odsc = strip_tags(odsc_m.group(1)) if odsc_m else ""
                        field["options"].append((olbl, odsc))

            elif re.search(r'<select', fc):
                field["type"] = "Keuzemenu (dropdown)"
                opts = re.findall(r'<option[^>]*value="([^"]*)"[^>]*>([^<]+)<', fc)
                field["options"] = [(o[1].strip(), "") for o in opts
                                    if o[0] not in ("", "--", "placeholder")]

            elif re.search(r'type="file"', fc):
                field["type"] = "Bestandsupload (PDF/JPG/PNG)"

            elif re.search(r'<textarea', fc):
                field["type"] = "Tekstveld (meerdere regels)"
                ph_m = re.search(r'placeholder="([^"]+)"', fc)
                if ph_m: field["placeholder"] = ph_m.group(1)

            elif re.search(r'type="date"', fc):
                field["type"] = "Datum"

            elif re.search(r'type="number"', fc):
                field["type"] = "Getal"
                ph_m = re.search(r'placeholder="([^"]+)"', fc)
                if ph_m: field["placeholder"] = ph_m.group(1)

            elif re.search(r'type="checkbox"', fc):
                field["type"] = "Bevestiging (checkbox)"

            elif re.search(r'<input', fc):
                field["type"] = "Tekstveld"
                ph_m = re.search(r'placeholder="([^"]+)"', fc)
                if ph_m: field["placeholder"] = ph_m.group(1)

            else:
                field["type"] = "Veld"

            fields.append(field)

        if fields:
            sections.append({"title": title, "fields": fields})

    return sections


def make_doc(form_name, sections):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    DARK  = RGBColor(10, 30, 66)
    GREY  = RGBColor(90, 100, 120)
    LIGHT = RGBColor(170, 178, 195)
    RULE  = RGBColor(210, 215, 225)

    # Title
    t = doc.add_heading(f"Vragenlijst – {form_name}", level=0)
    t.runs[0].font.color.rgb = DARK
    doc.add_paragraph(
        "Gebruik de kolom 'Nieuwe tekst / aanpassing' om de gewenste wijzigingen aan te geven."
    ).runs[0].font.color.rgb = GREY

    doc.add_paragraph()

    for sec in sections:
        h = doc.add_heading(sec["title"], level=1)
        for run in h.runs:
            run.font.color.rgb = DARK
            run.font.size = Pt(11)

        for field in sec["fields"]:
            # Question label
            p = doc.add_paragraph()
            r = p.add_run(field["label"])
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(15, 29, 58)

            # Type chip
            tp = doc.add_paragraph()
            tr = tp.add_run(f"Type: {field['type']}")
            tr.font.size = Pt(8.5)
            tr.font.color.rgb = LIGHT

            # Placeholder
            if field["placeholder"]:
                pp = doc.add_paragraph()
                pr = pp.add_run(f"Voorbeeld: {field['placeholder']}")
                pr.font.size = Pt(9)
                pr.font.color.rgb = GREY
                pr.italic = True

            # Hint
            if field["hint"]:
                hp = doc.add_paragraph()
                hr = hp.add_run(f"ℹ  {field['hint']}")
                hr.font.size = Pt(9)
                hr.font.color.rgb = GREY
                hr.italic = True

            # Extra text
            if field["extra_text"]:
                ep = doc.add_paragraph()
                er = ep.add_run(field["extra_text"][:200] + ("…" if len(field["extra_text"]) > 200 else ""))
                er.font.size = Pt(9)
                er.font.color.rgb = GREY
                er.italic = True

            # Options
            for (olbl, odsc) in field["options"]:
                op = doc.add_paragraph(style="List Bullet")
                op.paragraph_format.left_indent = Inches(0.25)
                ob = op.add_run(olbl)
                ob.font.size = Pt(10)
                if odsc:
                    od = op.add_run(f"  —  {odsc}")
                    od.font.size = Pt(9)
                    od.font.color.rgb = GREY

            # Revision line
            rev = doc.add_paragraph()
            rr = rev.add_run("Nieuwe tekst / aanpassing:  ___________________________________________")
            rr.font.size = Pt(9.5)
            rr.font.color.rgb = RULE

            doc.add_paragraph()

        doc.add_paragraph()

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────
for form_name, filename in FILES.items():
    path = PUBLIC_DIR / filename
    if not path.exists():
        print(f"SKIP: {filename}")
        continue

    print(f"Verwerken: {filename} ...")
    html = path.read_text(encoding="utf-8")
    sections = extract_form(html)

    if not sections:
        print(f"  ⚠ Geen velden gevonden")
        continue

    total = sum(len(s["fields"]) for s in sections)
    print(f"  {len(sections)} secties, {total} velden")

    doc  = make_doc(form_name, sections)
    safe = form_name.replace(" ", "_")
    out  = DESKTOP / f"Vragenlijst_{safe}.docx"
    doc.save(out)
    print(f"  ✓ Opgeslagen: {out.name}")

print("\nKlaar!")

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches

PUBLIC_DIR = Path(__file__).parent.parent / "public"
DESKTOP    = Path.home() / "Desktop"

FILES = {
    "BV Oprichten":           "vragenlijst-bv-oprichten.html",
    "BV en Holding Oprichten":"vragenlijst-bv-holding.html",
    "Holding Oprichten":      "vragenlijst-holding-oprichten.html",
    "Geruisloze Inbreng":     "vragenlijst-geruisloos.html",
    "Eenmanszaak naar BV":    "vragenlijst-eenmanszaak-naar-bv.html",
    "VOF naar BV":            "vragenlijst-vof-naar-bv.html",
}

def strip_tags(s):
    s = re.sub(r'<[^>]+>', ' ', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip(" *\n\r")

def extract_form(html):
    """Returns list of sections: [{title, fields:[{label,type,hint,options,option_descs,placeholder}]}]"""

    # Remove <style> and <script> blocks
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.S)
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.S)

    # Split on card divs
    # A card looks like: <div class="card"> ... <div class="card-title">X</div> ... fields ... </div>
    # We'll split the HTML on card-header openings to get sections
    sections = []

    card_blocks = re.split(r'(?=<div[^>]*class="[^"]*\bcard\b[^"]*")', html)

    for block in card_blocks:
        # Get section title from card-title
        title_m = re.search(r'<div[^>]*class="[^"]*card-title[^"]*"[^>]*>(.*?)</div>', block, re.S)
        if not title_m:
            continue
        title = strip_tags(title_m.group(1))

        fields = []

        # Find all <div class="field"...> blocks
        field_blocks = re.findall(
            r'<div[^>]*class="[^"]*\bfield\b[^"]*"[^>]*>(.*?)(?=<div[^>]*class="[^"]*\bfield\b[^"]*"|<div[^>]*class="[^"]*\bcard-footer\b[^"]*"|</div>\s*</div>\s*</div>\s*</div>)',
            block, re.S
        )

        for fb in field_blocks:
            # Label
            lbl_m = re.search(r'<label[^>]*class="[^"]*field-label[^"]*"[^>]*>(.*?)</label>', fb, re.S)
            if not lbl_m:
                continue
            label = strip_tags(lbl_m.group(1))
            if not label:
                continue

            field = {
                "label": label,
                "type": "",
                "hint": "",
                "options": [],
                "option_descs": [],
                "placeholder": "",
            }

            # Hint
            hint_m = re.search(r'<div[^>]*class="[^"]*field-hint[^"]*"[^>]*>(.*?)</div>', fb, re.S)
            if hint_m:
                field["hint"] = strip_tags(hint_m.group(1))

            # Determine type & extract options
            if re.search(r'class="[^"]*options-list[^"]*"', fb):
                # Radio / card-style options
                has_checkbox = bool(re.search(r'type="checkbox"', fb))
                field["type"] = "Meerkeuze (meerdere mogelijk)" if has_checkbox else "Keuze (één antwoord)"
                opt_labels = re.findall(r'<div[^>]*class="[^"]*option-label[^"]*"[^>]*>(.*?)</div>', fb, re.S)
                opt_descs  = re.findall(r'<div[^>]*class="[^"]*option-desc[^"]*"[^>]*>(.*?)</div>', fb, re.S)
                field["options"]      = [strip_tags(o) for o in opt_labels]
                field["option_descs"] = [strip_tags(d) for d in opt_descs]

            elif re.search(r'<select', fb):
                field["type"] = "Keuzemenu (dropdown)"
                opts = re.findall(r'<option[^>]*value="([^"]*)"[^>]*>([^<]+)<', fb)
                field["options"] = [o[1].strip() for o in opts if o[0] not in ("","--","placeholder")]

            elif re.search(r'<textarea', fb):
                field["type"] = "Tekstveld (meerdere regels)"
                ph_m = re.search(r'placeholder="([^"]+)"', fb)
                if ph_m: field["placeholder"] = ph_m.group(1)

            elif re.search(r'type="date"', fb):
                field["type"] = "Datum"

            elif re.search(r'type="number"', fb):
                field["type"] = "Getal"
                ph_m = re.search(r'placeholder="([^"]+)"', fb)
                if ph_m: field["placeholder"] = ph_m.group(1)

            elif re.search(r'type="file"', fb):
                field["type"] = "Bestandsupload"

            elif re.search(r'<input', fb):
                field["type"] = "Tekstveld"
                ph_m = re.search(r'placeholder="([^"]+)"', fb)
                if ph_m: field["placeholder"] = ph_m.group(1)

            else:
                field["type"] = "Veld"

            fields.append(field)

        if fields:
            sections.append({"title": title, "fields": fields})

    return sections


def make_doc(form_name, sections):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # Document title
    t = doc.add_heading(f"Vragenlijst – {form_name}", level=0)
    t.runs[0].font.color.rgb = RGBColor(10, 30, 66)
    doc.add_paragraph()

    BLUE  = RGBColor(10, 30, 66)
    GREY  = RGBColor(100, 110, 130)
    LIGHT = RGBColor(160, 170, 185)

    for sec in sections:
        # Section heading
        h = doc.add_heading(sec["title"], level=1)
        for run in h.runs:
            run.font.color.rgb = BLUE
            run.font.size = Pt(11)

        for field in sec["fields"]:
            # Question
            p = doc.add_paragraph()
            r = p.add_run(field["label"])
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(15, 29, 58)

            # Type
            tp = doc.add_paragraph()
            tr = tp.add_run(f"Type: {field['type']}")
            tr.font.size = Pt(8.5)
            tr.font.color.rgb = LIGHT

            # Placeholder
            if field["placeholder"]:
                pp = doc.add_paragraph()
                pr = pp.add_run(f"Voorbeeld: {field['placeholder']}")
                pr.font.size = Pt(9)
                pr.font.color.rgb = GREY
                pr.italic = True

            # Hint
            if field["hint"]:
                hp = doc.add_paragraph()
                hr = hp.add_run(f"Toelichting: {field['hint']}")
                hr.font.size = Pt(9)
                hr.font.color.rgb = GREY
                hr.italic = True

            # Options
            for i, opt in enumerate(field["options"]):
                op = doc.add_paragraph(style="List Bullet")
                op.paragraph_format.left_indent = Inches(0.25)
                or_ = op.add_run(opt)
                or_.font.size = Pt(10)
                # If there's a matching description, add it
                if i < len(field["option_descs"]) and field["option_descs"][i]:
                    or2 = op.add_run(f"  →  {field['option_descs'][i]}")
                    or2.font.size = Pt(9)
                    or2.font.color.rgb = GREY

            # Revision line
            rev = doc.add_paragraph()
            rr = rev.add_run("Nieuwe tekst / aanpassing: _________________________________________________")
            rr.font.size = Pt(9.5)
            rr.font.color.rgb = RGBColor(200, 205, 215)

            doc.add_paragraph()

        doc.add_paragraph()

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────
for form_name, filename in FILES.items():
    path = PUBLIC_DIR / filename
    if not path.exists():
        print(f"SKIP (niet gevonden): {filename}")
        continue

    print(f"Verwerken: {filename} ...")
    html = path.read_text(encoding="utf-8")
    sections = extract_form(html)

    if not sections:
        print(f"  ⚠ Geen velden gevonden")
        continue

    total = sum(len(s["fields"]) for s in sections)
    print(f"  {len(sections)} secties, {total} velden")

    doc  = make_doc(form_name, sections)
    safe = form_name.replace(" ", "_")
    out  = DESKTOP / f"Vragenlijst_{safe}.docx"
    doc.save(out)
    print(f"  ✓ Opgeslagen: {out.name}")

print("\nKlaar!")

import os, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html.parser import HTMLParser

# ── Config ────────────────────────────────────────────────────────────────────
PUBLIC_DIR = Path(__file__).parent.parent / "public"
DESKTOP    = Path.home() / "Desktop"

FILES = {
    "BV Oprichten":               "vragenlijst-bv-oprichten.html",
    "BV + Holding Oprichten":     "vragenlijst-bv-holding.html",
    "Holding Oprichten":          "vragenlijst-holding-oprichten.html",
    "Geruisloze Inbreng":         "vragenlijst-geruisloos.html",
    "Eenmanszaak naar BV":        "vragenlijst-eenmanszaak-naar-bv.html",
    "VOF naar BV":                "vragenlijst-vof-naar-bv.html",
}

# ── HTML text extractor ───────────────────────────────────────────────────────
class FormExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.items = []          # list of (kind, text)  kind: section|label|option|placeholder|help
        self._stack = []
        self._capture = False
        self._buf = ""
        self._in_script = False
        self._in_style  = False

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        cls   = attrs.get("class", "")
        self._stack.append(tag)

        if tag in ("script", "style"):
            self._in_script = tag == "script"
            self._in_style  = tag == "style"
            return

        # Section headers
        if tag in ("h2", "h3") and any(k in cls for k in ("intake-section", "section", "")):
            self._start_capture(f"section:{tag}")
        elif tag == "label":
            self._start_capture("label")
        elif tag == "option":
            val = attrs.get("value","")
            if val and val not in ("","--"):
                self._start_capture("option")
        elif tag == "textarea":
            ph = attrs.get("placeholder","")
            if ph:
                self.items.append(("placeholder", ph))
        elif tag == "input":
            ph = attrs.get("placeholder","")
            if ph and "••" not in ph:
                self.items.append(("placeholder", ph))
        elif tag == "p" and "help" in cls:
            self._start_capture("help")

    def handle_endtag(self, tag):
        if tag in ("script","style"):
            self._in_script = False
            self._in_style  = False
        if self._stack:
            self._stack.pop()
        if self._capture and self._capture.split(":")[0] in (tag,) or \
           (self._capture and self._capture == f"section:{tag}"):
            text = re.sub(r'\s+', ' ', self._buf).strip()
            if text and len(text) > 1:
                self.items.append((self._capture.split(":")[0], text))
            self._capture = False
            self._buf = ""

    def handle_data(self, data):
        if self._in_script or self._in_style:
            return
        if self._capture:
            self._buf += data

    def _start_capture(self, kind):
        self._capture = kind
        self._buf = ""


def clean(text):
    return re.sub(r'\s+', ' ', text).strip()


def extract_form(html_path):
    html = html_path.read_text(encoding="utf-8")

    sections = []
    current_section = {"title": None, "fields": []}

    # ── Extract section headings (h2/h3 inside intake-section divs) ──────────
    sec_pattern = re.compile(
        r'<div[^>]*class="[^"]*intake-section[^"]*"[^>]*>.*?<h[23][^>]*>(.*?)</h[23]>',
        re.S
    )
    # ── Extract all label texts ───────────────────────────────────────────────
    label_pattern = re.compile(r'<label[^>]*>(.*?)</label>', re.S)
    # ── Extract radio/checkbox options ───────────────────────────────────────
    option_pattern = re.compile(
        r'<label[^>]*class="[^"]*radio-option[^"]*"[^>]*>.*?</?(input)[^>]*>([^<]+)', re.S
    )
    # ── Extract select options ────────────────────────────────────────────────
    select_pattern = re.compile(r'<option[^>]*value="([^"]+)"[^>]*>([^<]+)</option>', re.S)

    def strip_tags(s):
        return re.sub(r'<[^>]+>', '', s).strip()
    def strip_req(s):
        # remove <span class="req">*</span>
        s = re.sub(r'<span[^>]*class="req"[^>]*>[^<]*</span>', '', s)
        return strip_tags(s).strip(" *")

    # Split HTML into logical blocks per section
    # Find all intake-section div positions
    sec_starts = [m.start() for m in re.finditer(r'<div[^>]*class="[^"]*intake-section[^"]*"', html)]
    sec_starts.append(len(html))

    result = []

    for i, start in enumerate(sec_starts[:-1]):
        end = sec_starts[i+1]
        chunk = html[start:end]

        # Section title
        title_m = re.search(r'<h[23][^>]*>(.*?)</h[23]>', chunk, re.S)
        title = strip_tags(title_m.group(1)) if title_m else "Algemeen"

        fields = []

        # Find all login-field divs after this section heading (in the chunk)
        # We look in html between this section start and next section start
        # Actually let's look at the HTML between end of this section's h-tag and next section
        # Simplification: scan all login-field divs in chunk
        field_blocks = re.findall(r'<div[^>]*class="[^"]*login-field[^"]*"[^>]*>(.*?)(?=<div[^>]*class="[^"]*(?:login-field|intake-section)[^"]*"|$)', chunk, re.S)

        for fb in field_blocks:
            # Label
            label_m = re.search(r'<label[^>]*>(.*?)</label>', fb, re.S)
            if not label_m:
                continue
            label_text = strip_req(label_m.group(1))
            if not label_text or label_text in ("*",):
                continue

            field = {"label": label_text, "type": None, "options": [], "placeholder": ""}

            # Determine type
            if re.search(r'<select', fb):
                field["type"] = "Keuzemenu"
                opts = re.findall(r'<option[^>]*value="([^"]+)"[^>]*>([^<]+)</option>', fb)
                field["options"] = [clean(o[1]) for o in opts if o[0] not in ("","--")]
            elif re.search(r'type="radio"', fb):
                field["type"] = "Keuze (radio)"
                opts = re.findall(r'<label[^>]*class="[^"]*radio-option[^"]*"[^>]*>(.*?)</label>', fb, re.S)
                field["options"] = [strip_tags(o).strip() for o in opts if strip_tags(o).strip()]
            elif re.search(r'type="checkbox"', fb):
                field["type"] = "Checkbox"
            elif re.search(r'<textarea', fb):
                field["type"] = "Tekstveld (meerdere regels)"
                ph_m = re.search(r'placeholder="([^"]+)"', fb)
                if ph_m: field["placeholder"] = ph_m.group(1)
            elif re.search(r'type="date"', fb):
                field["type"] = "Datum"
            elif re.search(r'type="number"', fb):
                field["type"] = "Getal"
            else:
                field["type"] = "Tekstveld"
                ph_m = re.search(r'placeholder="([^"]+)"', fb)
                if ph_m: field["placeholder"] = ph_m.group(1)

            fields.append(field)

        if fields:
            result.append({"title": title, "fields": fields})

    # If nothing found via intake-section, try a flat extraction
    if not result:
        field_blocks = re.findall(r'<div[^>]*class="[^"]*login-field[^"]*"[^>]*>(.*?)(?=<div[^>]*class="[^"]*login-field[^"]*"|</form>)', html, re.S)
        fields = []
        for fb in field_blocks:
            label_m = re.search(r'<label[^>]*>(.*?)</label>', fb, re.S)
            if not label_m: continue
            label_text = strip_req(label_m.group(1))
            if not label_text: continue
            fields.append({"label": label_text, "type": "Veld", "options": [], "placeholder": ""})
        if fields:
            result.append({"title": "Vragen", "fields": fields})

    return result


# ── Word document builder ─────────────────────────────────────────────────────
def make_doc(form_name, sections):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_heading(f"Vragenlijst – {form_name}", level=0)
    title.runs[0].font.color.rgb = RGBColor(10, 30, 66)

    doc.add_paragraph()

    for sec in sections:
        # Section heading
        h = doc.add_heading(sec["title"], level=1)
        h.runs[0].font.color.rgb = RGBColor(10, 30, 66)
        h.runs[0].font.size = Pt(12)

        for field in sec["fields"]:
            # Question label
            p = doc.add_paragraph()
            run = p.add_run(field["label"])
            run.bold = True
            run.font.size = Pt(10.5)

            # Type hint
            type_p = doc.add_paragraph()
            type_run = type_p.add_run(f"Type: {field['type']}")
            type_run.font.size = Pt(9)
            type_run.font.color.rgb = RGBColor(120, 130, 150)

            # Placeholder
            if field.get("placeholder"):
                ph_p = doc.add_paragraph()
                ph_run = ph_p.add_run(f"Voorbeeld: {field['placeholder']}")
                ph_run.font.size = Pt(9)
                ph_run.font.color.rgb = RGBColor(140, 150, 160)
                ph_run.font.italic = True

            # Options
            if field["options"]:
                for opt in field["options"]:
                    op = doc.add_paragraph(style="List Bullet")
                    op.add_run(opt).font.size = Pt(10)

            # Revision space
            rev = doc.add_paragraph()
            rev_run = rev.add_run("Nieuwe tekst: _______________________________________________")
            rev_run.font.size = Pt(10)
            rev_run.font.color.rgb = RGBColor(180, 180, 180)

            doc.add_paragraph()  # spacing

        doc.add_paragraph()

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────
for form_name, filename in FILES.items():
    path = PUBLIC_DIR / filename
    if not path.exists():
        print(f"SKIP (not found): {filename}")
        continue

    print(f"Verwerken: {filename} ...")
    sections = extract_form(path)

    if not sections:
        print(f"  ⚠ Geen velden gevonden in {filename}")
        continue

    total = sum(len(s["fields"]) for s in sections)
    print(f"  {len(sections)} secties, {total} velden")

    doc  = make_doc(form_name, sections)
    safe = form_name.replace(" ", "_").replace("+","en")
    out  = DESKTOP / f"Vragenlijst_{safe}.docx"
    doc.save(out)
    print(f"  ✓ Opgeslagen: {out.name}")

print("\nKlaar!")
